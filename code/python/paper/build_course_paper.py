from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("write_paper") / "医疗体制改革与县域药店进入_面板数据课程论文.docx"


def set_east_asian_font(run, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10)
    set_east_asian_font(r)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_paragraph(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_east_asian_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120) if level >= 2 else RGBColor(46, 116, 181)
    set_east_asian_font(r, "黑体")
    return p


def add_table(doc, title, headers, rows):
    caption = add_paragraph(doc, title)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        shade_cell(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    add_paragraph(doc, "注：括号内为按区县聚类的稳健标准误。*、**、*** 分别表示在 5%、1%、0.1% 水平显著。")
    return table


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        s = styles[style_name]
        s.font.name = "黑体"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.paragraph_format.line_spacing = 2
        s.paragraph_format.space_before = Pt(8)
        s.paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    run = title.add_run("医疗体制改革与县域药店进入")
    run.bold = True
    run.font.size = Pt(18)
    set_east_asian_font(run, "黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = None
    r = subtitle.add_run("基于2012—2019年区县面板数据的实证分析")
    r.font.size = Pt(14)
    set_east_asian_font(r, "黑体")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.first_line_indent = None
    r = meta.add_run("面板数据计量经济学专题课程结课报告")
    r.font.size = Pt(11)
    set_east_asian_font(r)

    add_heading(doc, "摘要", 1)
    add_paragraph(
        doc,
        "本文围绕医疗体制改革对县域药店进入的影响展开实证研究。研究使用2012—2019年中国区县—年份面板数据，"
        "将各区县政策年份的下一年定义为改革实际生效年份，并以区县药店数量、药店数量对数、ln(药店数量+1)以及是否存在药店作为主要被解释变量。"
        "识别策略首先采用包含区县固定效应和年份固定效应的交错双重差分模型，并控制人口规模、夜间灯光强度和道路密度；随后使用Callaway和Sant'Anna"
        "的CSDID方法、PPML计数模型、平衡样本和省份—年份固定效应进行稳健性检验，同时构造以基准公立医院强度为处理强度的连续DID及工具变量估计。"
        "基准TWFE结果显示，改革后区县药店数量平均增加3.739家，药店数量对数增加0.031，ln(药店数量+1)增加0.128，是否有药店的概率提高4.7个百分点。"
        "但CSDID结果显示平均处理效应并不显著，且数量结果存在较强前趋势，说明简单TWFE估计可能混合了不同改革批次的异质趋势。机制分析表明，改革后的增量主要体现为连锁药店数量上升，非连锁药店变化不显著。"
        "总体而言，医疗体制改革与县域药店进入之间存在正相关证据，但严格因果解释仍需谨慎。"
    )
    kw = add_paragraph(doc, "关键词：医疗体制改革；药店进入；区县面板数据；交错双重差分；CSDID")
    kw.runs[0].bold = True

    add_heading(doc, "一、引言", 1)
    intro = [
        "医疗服务体系改革不仅影响医院内部的服务供给、药品采购和医生行为，也可能改变医院周边乃至县域范围内的药品零售市场结构。"
        "在传统医疗体制下，患者购药需求与医院处方、医保支付、药品流通渠道紧密相连。若改革削弱医院对药品销售的依赖，改善处方外流环境，或提高居民在基层医疗体系中的就诊频率，县域药店可能获得更多进入机会。"
        "相反，如果改革使医院药事服务更集中、药品配送更规范，零售药店面对的竞争压力也可能上升。因此，医疗体制改革对药店数量的影响并非理论上显然为正或为负，需要借助面板数据和政策时点差异进行实证检验。",
        "从计量经济学角度看，这一问题适合使用面板数据方法。第一，不同区县在改革推进时间上存在差异，形成了交错实施的政策环境；第二，区县药店数量具有明显的地区固定特征，例如人口规模、城市化水平、商业环境和原有医疗资源，这些难以完全观测的因素可以通过区县固定效应吸收；第三，药店市场又受到全国性宏观趋势、监管环境、药品零售连锁化和移动互联网医疗发展的共同影响，这些年份层面的共同冲击可以由年份固定效应控制。",
        "本文的问题是：县域医疗体制改革是否促进了药店进入？如果促进，其影响主要体现在药店数量水平、相对增长还是进入边际？影响是否由连锁药店推动？在交错政策实施背景下，传统TWFE估计是否稳健？围绕这些问题，本文基于已有区县层面结果撰写课程论文，重点展示面板模型的构建、估计和解释。"
    ]
    for paragraph in intro:
        add_paragraph(doc, paragraph)

    add_heading(doc, "二、数据来源与变量构造", 1)
    data_text = [
        "本文使用的核心数据为区县药店面板数据 `county_pharmacy_panel_all_years.dta`，并合并区县层面的控制变量数据 `county_panel_1020.dta`。"
        "药店面板包含2012—2022年区县药店数量及连锁、非连锁药店信息；由于道路密度等主要控制变量在当前清洗结果中主要覆盖至2019年，基准回归样本限定在2012—2019年。"
        "合并后的基准样本包含17463个区县—年份观测，标准误按区县聚类。",
        "被解释变量包括四类。第一，`pharmacy_count` 表示区县当年药店存量数量，是最直接的进入结果。第二，`ln_pharmacy` 为药店数量取自然对数，零药店区县在该指标中缺失，因此样本数下降至15940。"
        "第三，`ln_pharmacy_plus1=ln(pharmacy_count+1)` 保留零药店观测，用于检验对数规格是否受零值处理影响。第四，`has_pharmacy` 表示区县是否至少存在一家药店，刻画扩展边际。",
        "核心解释变量为 `did`，即区县进入改革后的虚拟变量。已有代码将政策年份的下一年设为实际生效年份，定义为 `effective_policy_year=policy_year+1`，随后令 `did=1{year>=effective_policy_year}`。"
        "这一处理考虑到制度改革从政策发布到市场主体反应之间可能存在执行和调整滞后。控制变量包括 `ln_pop`、`ln_nl_mean` 和 `ln_road_den`，分别代表人口规模、夜间灯光强度和道路密度，用于控制县域需求规模、经济活跃度和交通可达性。",
        "描述统计显示，基准样本中区县平均药店数量为97.12家，标准差为109.80，中位数为64，最大值达到1396，说明药店分布存在明显右偏。约91.3%的区县—年份至少有一家药店。连锁药店平均数量为8.08家，非连锁药店平均数量为89.04家，表明样本期内非连锁药店仍占主体，但连锁药店可能是改革后更敏感的市场进入者。"
    ]
    for paragraph in data_text:
        add_paragraph(doc, paragraph)

    add_table(
        doc,
        "表1 主要变量描述统计",
        ["变量", "样本量", "均值", "标准差", "中位数"],
        [
            ["药店数量", "17507", "97.122", "109.799", "64"],
            ["ln(药店数量)", "15984", "4.157", "1.155", "4.290"],
            ["ln(药店数量+1)", "17507", "3.829", "1.575", "4.174"],
            ["是否有药店", "17507", "0.913", "0.282", "1"],
            ["连锁药店数量", "17507", "8.085", "27.110", "0"],
            ["非连锁药店数量", "17507", "89.037", "97.882", "60"],
        ],
    )

    add_heading(doc, "三、模型设定与识别思路", 1)
    model_text = [
        "基准模型为交错双重差分的双向固定效应模型：",
        "Y_ct = beta DID_ct + X_ct'gamma + mu_c + lambda_t + epsilon_ct。",
        "其中，Y_ct 为区县 c 在年份 t 的药店结果变量；DID_ct 表示该区县在年份 t 是否已进入改革后；X_ct 包括人口、夜间灯光和道路密度；mu_c 为区县固定效应，lambda_t 为年份固定效应。"
        "系数 beta 是本文最关注的参数，可解释为在控制区县不随时间变化特征、年份共同冲击和可观测县域发展条件后，改革后区县药店结果相对于改革前及尚未改革区县的平均变化。",
        "这一模型的核心识别假设是，在没有医疗体制改革的情况下，不同改革批次区县的药店进入趋势应当相近。由于政策实施时间交错，传统TWFE在存在异质处理效应时可能把已处理组错误地作为其他处理组的对照，从而导致加权解释复杂。"
        "因此，本文将TWFE作为基准结果，同时引入CSDID估计组别—时间平均处理效应，并进行事件研究和前趋势检验。",
        "此外，药店数量是非负计数变量，且存在零值。为检验线性模型对计数结果的适用性，本文使用PPML模型估计改革对药店数量的比例影响。为探索改革强度差异，本文还使用基准公立医院数量的对数构造连续处理强度，即 `post_policy × ln(public_hospital_base+1)`，并使用1959年流行病防治站存量或1950年代防治站年数与改革后变量交互作为工具变量。这一设计尝试利用历史公共卫生资源分布解释基准公立医院强度。"
    ]
    for paragraph in model_text:
        add_paragraph(doc, paragraph)

    add_heading(doc, "四、基准回归结果", 1)
    results_text = [
        "表2报告基准TWFE结果。以药店数量为因变量时，未加入控制变量的固定效应模型估计系数为2.143，加入人口、夜间灯光和道路密度后系数上升为3.739，并在5%水平显著。"
        "这意味着改革后区县平均药店数量增加约3.7家。相对于样本均值97.1家，这一水平效应并不巨大，但在县域年度变化中具有实际含义。",
        "在对数规格中，ln(药店数量+1)的估计系数为0.128，且在0.1%水平显著，说明在保留零药店区县后，改革与药店数量相对增长显著相关。"
        "ln(药店数量)的系数为0.031，在5%水平显著，表示药店数量约增加3.1%。是否有药店的线性概率模型系数为0.047，说明改革后区县出现至少一家药店的概率提高约4.7个百分点。"
        "这些结果整体支持改革促进药店进入的初步判断。",
        "控制变量的方向也具有经济含义。人口规模在药店数量模型中系数较大且显著，反映需求规模越大的县域药店越多；夜间灯光强度和道路密度在数量模型中也为正，说明经济活跃度和交通可达性有助于零售药店布局。"
        "不过，在对数和进入边际模型中部分控制变量方向不同，提示县域规模、药店存量和是否已有药店之间并非简单线性关系。"
    ]
    for paragraph in results_text:
        add_paragraph(doc, paragraph)

    add_table(
        doc,
        "表2 基准TWFE估计结果",
        ["因变量", "DID系数", "标准误", "样本量", "R²"],
        [
            ["药店数量", "3.739*", "(1.565)", "17463", "0.842"],
            ["ln(药店数量+1)", "0.128***", "(0.020)", "17463", "0.872"],
            ["ln(药店数量)", "0.031*", "(0.014)", "15940", "0.876"],
            ["是否有药店", "0.047***", "(0.005)", "17463", "0.663"],
        ],
    )

    add_heading(doc, "五、稳健性检验与动态效应", 1)
    robust_text = [
        "首先，PPML计数模型得到的DID系数为0.053，标准误为0.013，且在0.1%水平显著。将其近似换算为比例影响，改革后药店数量约增加5.4%。"
        "PPML保留零值并适用于非负计数结果，这一结果说明基准TWFE的正向结论并非完全由线性数量模型驱动。",
        "其次，在限制为中间改革批次、平衡面板以及加入省份—年份固定效应后，结果出现差异。中间批次样本中，药店数量系数为3.374但未达到5%显著水平，ln(药店数量+1)系数为0.177且显著；平衡样本中数量系数为2.993，ln(药店数量+1)系数为0.117且显著。"
        "加入省份—年份固定效应后，数量系数降至0.916，ln(药店数量+1)系数为0.027且不显著。这表明一部分基准效应可能来自省域层面的同步趋势或地区政策环境变化。",
        "第三，CSDID结果对TWFE结论形成重要约束。以药店数量为结果时，CSDID的事后平均效应为4.871，标准误为3.138，p值为0.121，未达到传统显著性水平；以ln(药店数量)为结果时，事后平均效应为-0.0277，标准误为0.0690，亦不显著。"
        "更重要的是，药店数量的前趋势检验拒绝原假设，chi2(10)=67.373，p值为0.0000；对数结果的前趋势检验p值为0.0633，接近10%显著性水平。"
        "这说明不同改革批次在改革前的药店数量变化并不完全平行，因而基准TWFE的因果解释需要谨慎。",
        "事件研究图和五类估计量比较进一步显示，改革后效应并非在所有估计方法中一致。TWFE给出小幅正向平均结果，而CSDID静态平均效应不显著。"
        "在交错DID环境下，这种差异并不罕见，因为TWFE估计量会受到不同处理时点和处理效应异质性的影响。本文因此将TWFE结果解释为相关性较强的基准证据，而不是无需保留的最终因果结论。"
    ]
    for paragraph in robust_text:
        add_paragraph(doc, paragraph)

    add_table(
        doc,
        "表3 稳健性与替代估计结果",
        ["模型或样本", "药店数量效应", "对数效应", "说明"],
        [
            ["PPML", "0.053***", "-", "约对应5.4%的数量增加"],
            ["中间改革批次", "3.374", "0.177***", "有效政策年2015—2017"],
            ["平衡面板", "2.993", "0.117***", "每县8年均有观测"],
            ["省份—年份固定效应", "0.916", "0.027", "吸收省域年度冲击"],
            ["CSDID", "4.871", "-0.028", "事后平均效应均不显著"],
        ],
    )

    add_heading(doc, "六、机制分析：连锁化、空间距离与公立医院强度", 1)
    mech_text = [
        "从药店类型看，改革后的药店数量增加主要来自连锁药店。连锁药店数量的DID系数为4.417，标准误为0.459，显著为正；非连锁药店系数为-0.678且不显著；连锁药店份额系数为0.013，显著为正。"
        "这意味着医疗体制改革可能并不是简单增加所有药店，而是更有利于组织化、标准化程度较高的连锁药店进入。一个可能解释是，连锁药店在处方承接、医保合规、供应链管理和规模化采购方面更具优势，因而更能响应改革后出现的市场机会。",
        "从空间距离看，将药店按是否位于距离最近医院2公里以内划分后，2公里以内药店数量的改革系数为1.267且不显著，2公里以外药店数量的系数为2.471并在5%水平显著。"
        "这一结果提示，县域药店增长不一定只发生在医院周边，也可能表现为更广泛的县域零售网络扩张。若改革提升居民基层就医或院外购药需求，药店进入可能向社区、乡镇和交通节点扩散，而不仅仅围绕医院门口集中。",
        "连续DID进一步显示，改革效应与基准公立医院强度有关。以 `post_policy × ln(public_hospital_base+1)` 为核心解释变量时，药店数量系数为20.308，标准误为3.460，显著为正；但ln(药店数量)系数为-0.028且不显著。"
        "使用1959年流行病防治站存量作为工具变量后，数量效应扩大为79.687，弱工具检验统计量为100.124，说明一阶段强度较高；但对数效应仍不显著。另一个以1950年代防治站年数为工具的估计也得到显著正向数量效应。"
        "这些结果表明，医疗资源基础较强的县域在改革后可能吸引更多药店数量进入，但这种效应更表现为水平数量增加，而不是稳定的比例增长。"
    ]
    for paragraph in mech_text:
        add_paragraph(doc, paragraph)

    doc.add_page_break()
    add_table(
        doc,
        "表4 机制与异质性结果",
        ["结果变量或模型", "核心系数", "标准误", "解释"],
        [
            ["连锁药店数量", "4.417***", "(0.459)", "改革后连锁药店显著增加"],
            ["非连锁药店数量", "-0.678", "(1.394)", "非连锁药店变化不显著"],
            ["连锁药店份额", "0.013***", "(0.002)", "连锁化程度提高"],
            ["2公里以内药店", "1.267", "(0.972)", "医院近邻药店增加不显著"],
            ["2公里以外药店", "2.471*", "(0.990)", "更远范围药店显著增加"],
            ["连续DID：数量", "20.308***", "(3.460)", "公立医院基准强度越高，数量效应越大"],
            ["IV连续DID：数量", "79.687***", "(13.734)", "以1959年防治站存量为工具"],
        ],
    )

    add_heading(doc, "七、讨论：识别局限与经济解释", 1)
    discussion = [
        "本文结果可以从需求、供给和制度环境三个角度理解。需求侧，医疗体制改革可能改变居民就医与购药路径，扩大院外购药需求。供给侧，改革可能降低药店进入县域市场的不确定性，尤其是连锁药店能够通过更规范的供应链和医保对接能力快速扩张。"
        "制度侧，改革往往伴随基层医疗服务、药品采购和医保支付方式调整，地方政府在医疗服务体系建设中的行动也可能带来药店市场的同步变化。",
        "然而，本文的识别仍存在局限。首先，CSDID和事件研究显示数量结果存在明显前趋势，这意味着改革前不同批次区县已处于不同药店增长轨道。若较早改革的区县本身就是药店扩张更快或医疗资源增长更快的地区，TWFE估计可能高估改革效应。"
        "其次，控制变量虽包括人口、夜间灯光和道路密度，但无法完全吸收医保政策、药品监管、地方财政能力、电商药品销售和连锁企业战略布局等时间变化因素。第三，药店数量是存量指标，新增进入、退出和并购整合无法被完全区分，连锁药店增加也可能包含收购改牌而非净新增门店。",
        "因此，本文更稳妥的结论是：在区县层面，医疗体制改革与药店市场扩张存在正向关联，且这种关联主要体现为连锁药店和数量水平的增加；但若要将其解释为严格的平均因果效应，需要更强的平行趋势证据或更精细的识别设计。"
        "后续研究可以尝试使用更短的政策窗口、匹配相近改革批次、引入县域线性趋势、使用行政边界附近比较，或构建门店级进入退出数据，以提高识别可信度。"
    ]
    for paragraph in discussion:
        add_paragraph(doc, paragraph)

    add_heading(doc, "八、结论", 1)
    conclusion = [
        "本文基于2012—2019年中国区县—年份面板数据，研究医疗体制改革对县域药店进入的影响。基准TWFE结果显示，改革后区县药店数量、对数数量和是否存在药店均出现上升，PPML模型也支持约5.4%的正向比例效应。"
        "机制分析表明，增量主要来自连锁药店，且2公里以外药店数量的增加更显著，说明改革可能推动县域药品零售网络扩张，而不仅是医院周边药店集聚。",
        "同时，本文也发现稳健性结果并不完全支持强因果结论。CSDID估计的事后平均效应不显著，药店数量结果存在明显改革前趋势，加入省份—年份固定效应后基准效应明显减弱。"
        "因此，本文最终认为：医疗体制改革与县域药店进入存在较稳定的正相关证据，尤其体现在连锁药店扩张和药店数量水平增加上；但由于交错改革批次的前趋势和地区异质性问题，不能简单将TWFE系数解释为完全可信的平均处理效应。",
        "从面板数据计量方法的角度看，本文展示了双向固定效应模型在政策评估中的直观优势，也说明在交错DID背景下必须结合CSDID、事件研究、PPML和异质性分析进行交叉验证。对于医疗改革这类复杂政策，研究结论不应只依赖单一模型的显著性，而应综合考虑识别假设、动态趋势和经济机制。"
    ]
    for paragraph in conclusion:
        add_paragraph(doc, paragraph)

    add_heading(doc, "参考文献", 1)
    refs = [
        "Callaway, B., & Sant'Anna, P. H. C. (2021). Difference-in-Differences with multiple time periods. Journal of Econometrics, 225(2), 200-230.",
        "Sun, L., & Abraham, S. (2021). Estimating dynamic treatment effects in event studies with heterogeneous treatment effects. Journal of Econometrics, 225(2), 175-199.",
        "Borusyak, K., Jaravel, X., & Spiess, J. (2021). Revisiting event study designs: Robust and efficient estimation. Working paper.",
        "Cengiz, D., Dube, A., Lindner, A., & Zipperer, B. (2019). The effect of minimum wages on low-wage jobs. Quarterly Journal of Economics, 134(3), 1405-1454.",
        "Wooldridge, J. M. (2010). Econometric Analysis of Cross Section and Panel Data. MIT Press.",
        "已有区县层面回归结果：`D:\\Hospital_Pharmacy program\\tables\\regression\\integrated_results\\county_v2`。"
    ]
    for ref in refs:
        p = add_paragraph(doc, ref)
        p.paragraph_format.first_line_indent = None

    # Add page numbers in the footer.
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build_doc()
